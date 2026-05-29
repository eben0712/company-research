"""Markdown → Word (.docx) 导出脚本
用法: python export_to_docx.py <input.md> [output.docx]
如不指定输出路径，默认与输入同名的.docx文件。
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 配色方案 ──────────────────────────────────────────
COLOR_PRIMARY   = "1F3864"   # 深蓝 — 标题、表头
COLOR_ACCENT    = "2E75B6"   # 中蓝 — 二级标题
COLOR_BODY      = "000000"   # 纯黑 — 正文
COLOR_TABLE_BG  = "F2F6FC"   # 浅蓝灰 — 表格交替行
COLOR_CODE_BG   = "F5F5F5"   # 浅灰 — 代码块背景
COLOR_BORDER    = "B0B0B0"   # 边框色
COLOR_SEPARATOR = "CCCCCC"   # 分隔线

FONT_BODY       = "宋体"
FONT_HEADING    = "黑体"
FONT_CODE       = "Consolas"
SIZE_BODY       = 10.5
SIZE_SMALL      = 9
SIZE_TITLE      = 18
SIZE_H1         = 14
SIZE_H2         = 12
SIZE_H3         = 12

# ── 辅助函数 ──────────────────────────────────────────

def _set_font(run, name=FONT_BODY, size=SIZE_BODY, color=COLOR_BODY, bold=False):
    """给一个 run 设置字体属性（中英文统一）。"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def _set_cell_bg(cell, color):
    """给表格单元格设置背景色。"""
    tcPr = cell._element.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    tcPr.append(shading)


def _set_para_spacing(paragraph, before=0, after=0, line_spacing=1.35):
    """设置段落间距。"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_page_numbers(doc):
    """在页脚添加页码。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=6, after=0)

        run1 = p.add_run('— ')
        _set_font(run1, size=8, color="999999")
        # PAGE field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_page = p.add_run()
        run_page._element.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_page._element.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_page._element.append(fldChar2)

        run2 = p.add_run(' —')
        _set_font(run2, size=8, color="999999")


def _add_top_border(paragraph, color=COLOR_PRIMARY, width=8):
    """给段落加顶部边框线。"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{width}" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ── 文档设置 ──────────────────────────────────────────

def setup_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.4)
        section.bottom_margin = Cm(1.7)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)

    # Normal style — 宋体 10.5pt，1.5 倍行距
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(SIZE_BODY)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_BODY)
    rFonts.set(qn('w:ascii'), FONT_BODY)

    # Heading styles — 黑体，蓝色
    for lvl, (size, color) in enumerate([
        (SIZE_TITLE, COLOR_PRIMARY),
        (SIZE_H1, COLOR_PRIMARY),
        (SIZE_H2, COLOR_ACCENT),
        (SIZE_H3, COLOR_PRIMARY),
    ]):
        hstyle = doc.styles[f'Heading {lvl + 1}'] if lvl > 0 else doc.styles['Title']
        hstyle.font.name = FONT_HEADING
        hstyle.font.size = Pt(size)
        hstyle.font.color.rgb = RGBColor.from_string(color)
        hstyle.font.bold = True
        hstyle.paragraph_format.line_spacing = 1.3
        hstyle.paragraph_format.space_before = Pt(0)
        hstyle.paragraph_format.space_after = Pt(0)
        rpr = hstyle.element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        rFonts.set(qn('w:ascii'), FONT_HEADING)

    _add_page_numbers(doc)
    return doc


# ── 内容元素 ──────────────────────────────────────────

def add_title(doc, text):
    """主标题 —— 居中、大号、深蓝。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=48, after=30, line_spacing=1.3)
    run = p.add_run(text)
    _set_font(run, name=FONT_HEADING, size=SIZE_TITLE, color=COLOR_PRIMARY, bold=True)


def add_heading_custom(doc, text, level):
    """各级标题，黑体 + 蓝色，间距对标锅圈报告。"""
    if level == 0:
        add_title(doc, text)
        return

    sizes   = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3}
    colors  = {1: COLOR_PRIMARY, 2: COLOR_ACCENT, 3: COLOR_PRIMARY}
    before  = {1: 24, 2: 18, 3: 12}   # 段前（pt）
    after   = {1: 10, 2: 8,  3: 4}    # 段后（pt）

    size  = sizes.get(level, SIZE_H3)
    color = colors.get(level, COLOR_PRIMARY)
    sb    = before.get(level, 6)
    sa    = after.get(level, 3)

    p = doc.add_paragraph()
    _set_para_spacing(p, before=sb, after=sa, line_spacing=1.3)
    run = p.add_run(text)
    _set_font(run, name=FONT_HEADING, size=size, color=color, bold=True)

    # H1 加顶部边框
    if level == 1:
        _add_top_border(p, color=COLOR_PRIMARY, width=8)


def add_formatted_para(doc, text, font_name=FONT_BODY, size=SIZE_BODY, color=COLOR_BODY):
    """正文段落，支持 **粗体** 和 > 引用。"""
    p = doc.add_paragraph()
    _set_para_spacing(p, before=1, after=1, line_spacing=1.5)

    # 引用样式处理
    if text.startswith('> '):
        text = text[2:]
        # 加左侧缩进 + 灰色边框
        pf = p.paragraph_format
        pf.left_indent = Cm(0.8)
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="12" w:space="8" w:color="{COLOR_ACCENT}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        color = "555555"

    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            _set_font(run, name=font_name, size=size, color=color, bold=True)
        else:
            run = p.add_run(part)
            _set_font(run, name=font_name, size=size, color=color)
    return p


def add_code_para(doc, text):
    """代码块行 —— 灰底 + 缩进 + 等宽字体。"""
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=0, line_spacing=1.2)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    # 给段落加背景色（用 shading）
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_CODE_BG}" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    _set_font(run, name=FONT_CODE, size=9, color="444444")
    return p


def add_separator(doc):
    """水平分隔线。"""
    p = doc.add_paragraph()
    _set_para_spacing(p, before=12, after=12, line_spacing=1.0)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="2" w:space="1" w:color="{COLOR_SEPARATOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ── 表格 ──────────────────────────────────────────────

def add_table_from_markdown(doc, header_cells, data_rows):
    """专业表格：深蓝表头白字 + 交替行底色 + 自动列宽适配。"""
    ncols = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # ── 表头行 ──
    for i, h in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=2, after=2, line_spacing=1.1)
        run = p.add_run(h.strip())
        _set_font(run, size=SIZE_SMALL, color="FFFFFF", bold=True)
        _set_cell_bg(cell, COLOR_PRIMARY)

    # ── 表头行高 ──
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="400" w:hRule="atLeast"/>')
    trPr.append(trHeight)

    # ── 数据行 ──
    for r, row in enumerate(data_rows):
        bg = COLOR_TABLE_BG if r % 2 == 0 else "FFFFFF"
        for c in range(ncols):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            _set_para_spacing(p, before=1, after=1, line_spacing=1.1)
            val = row[c].strip() if c < len(row) else ''
            run = p.add_run(val)
            _set_font(run, size=SIZE_SMALL, color=COLOR_BODY)
            _set_cell_bg(cell, bg)

    # ── 列宽：简单按内容量比例分配 ──
    total_chars = [0] * ncols
    all_rows = [header_cells] + data_rows
    for row in all_rows:
        for c in range(min(ncols, len(row))):
            total_chars[c] = max(total_chars[c], len(row[c].strip()) if isinstance(row[c], str) else len(str(row[c])))

    char_total = sum(total_chars) or 1
    # 可用宽度约 18cm（A4 21cm - 2×1.5cm margins）
    avail = 18.0
    for c in range(ncols):
        ratio = total_chars[c] / char_total
        width = max(2.0, avail * ratio)
        # 第一列通常更宽（标签列），给个最小值
        if c == 0:
            width = max(2.5, width)
        for row in table.rows:
            row.cells[c].width = Cm(width)

    return table


def _flush_table_buffer(doc, table_buffer):
    if len(table_buffer) >= 2:
        header = [c.strip() for c in table_buffer[0].split('|')[1:-1]]
        data = []
        for tl in table_buffer[2:]:
            cells = [c.strip() for c in tl.split('|')[1:-1]]
            if any(c for c in cells):
                data.append(cells)
        if header and data:
            add_table_from_markdown(doc, header, data)


# ── 主转换逻辑 ────────────────────────────────────────

def convert(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = setup_document()
    table_buffer = []
    in_code_block = False

    for line in lines:
        stripped = line.rstrip()

        # ── 代码块 ──
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            if table_buffer:
                _flush_table_buffer(doc, table_buffer)
                table_buffer = []
            continue
        if in_code_block:
            add_code_para(doc, '    ' + stripped)
            continue

        # ── 空行：刷出待处理表格 ──
        if not stripped:
            if table_buffer:
                _flush_table_buffer(doc, table_buffer)
                table_buffer = []
            continue

        # ── 表格行 ──
        if stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(stripped)
            continue

        # ── 表格后第一条非表格行：先刷出表格 ──
        if table_buffer:
            _flush_table_buffer(doc, table_buffer)
            table_buffer = []

        # ── 标题 ──
        if stripped.startswith('# '):
            add_heading_custom(doc, stripped[2:], level=0)
        elif stripped.startswith('## '):
            add_heading_custom(doc, stripped[3:], level=1)
        elif stripped.startswith('### '):
            add_heading_custom(doc, stripped[4:], level=2)
        elif stripped.startswith('#### '):
            add_heading_custom(doc, stripped[5:], level=3)

        # ── 分隔线 ──
        elif stripped.strip().replace('-', '') == '' and len(stripped) > 5:
            add_separator(doc)

        # ── 列表项 ──
        elif stripped.startswith('- '):
            add_formatted_para(doc, '  ' + stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            add_formatted_para(doc, '  ' + text)

        # ── 正文 ──
        else:
            add_formatted_para(doc, stripped)

    # ── 文件末尾刷出 ──
    if table_buffer:
        _flush_table_buffer(doc, table_buffer)

    doc.save(output_path)
    print(f'Word saved: {output_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python export_to_docx.py <input.md> [output.docx]')
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else input_path.with_suffix('.docx')
    convert(str(input_path), str(output_path))
